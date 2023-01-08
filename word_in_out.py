from os.path import join, dirname, realpath
from os import chdir, listdir

import docx


class DocxReadWrite:

    def __init__(self, name_file_write, name_directory_read):
        self.name_file_write = name_file_write
        self.name_directory_read = name_directory_read


    def word_read(self) -> list:
        base_path = self.name_directory_read
        chdir(join(dirname(realpath(__file__)), base_path))
        file_in = listdir(path=".")[0]
        document = docx.Document(file_in)
        pages = document.paragraphs
        text_temp: list = [paragraph.text for paragraph in pages]
        return text_temp

    def word_write(self, text_out: list) -> None:
        document = docx.Document()
        for var in text_out:
            add_par = document.add_paragraph(var)
            fmt = add_par.paragraph_format
            fmt.space_before = 0
            fmt.space_after= 0
        document.save(self.name_file_write)